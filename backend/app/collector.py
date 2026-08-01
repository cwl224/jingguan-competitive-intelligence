from __future__ import annotations

import asyncio
import csv
import hashlib
import ipaddress
import io
import json
import os
import re
import socket
import threading
import time
from dataclasses import dataclass, field
from datetime import UTC, datetime
from email.utils import parsedate_to_datetime
from pathlib import Path
from typing import Any
from urllib.parse import urljoin, urlparse
from urllib.robotparser import RobotFileParser

import feedparser
import httpx
from bs4 import BeautifulSoup

from .config import Settings
from .database import Database
from .processing import DocumentProcessor


PARSER_VERSION = "collector-0.2.0"
ONLINE_SOURCE_TYPES = {
    "webpage",
    "dynamic_webpage",
    "sitemap",
    "rss",
    "public_api",
    "social_api",
    "public_database",
}
API_SOURCE_TYPES = {"public_api", "social_api", "public_database"}


class CollectionError(RuntimeError):
    def __init__(self, code: str, message: str, *, retryable: bool = False):
        super().__init__(message)
        self.code = code
        self.retryable = retryable


@dataclass(slots=True)
class FetchedResponse:
    url: str
    status_code: int
    content: bytes
    content_type: str
    headers: dict[str, str]


@dataclass(slots=True)
class CollectedArtifact:
    canonical_url: str
    title: str
    raw_content: bytes
    readable_text: str
    content_type: str
    published_at: str | None = None
    language: str | None = None
    metadata: dict[str, Any] = field(default_factory=dict)
    structured_fields: dict[str, Any] = field(default_factory=dict)


def _clean_text(value: str) -> str:
    lines = [re.sub(r"\s+", " ", line).strip() for line in value.splitlines()]
    return "\n".join(line for line in lines if line)


def _decode_bytes(content: bytes, content_type: str = "") -> str:
    charset_match = re.search(r"charset=([\w.-]+)", content_type, re.I)
    encodings = [charset_match.group(1)] if charset_match else []
    encodings.extend(["utf-8-sig", "utf-8", "gb18030"])
    for encoding in encodings:
        try:
            return content.decode(encoding)
        except (LookupError, UnicodeDecodeError):
            continue
    return content.decode("utf-8", errors="replace")


def _normalise_datetime(value: Any) -> str | None:
    if value is None or value == "":
        return None
    if isinstance(value, datetime):
        parsed = value
    else:
        text = str(value).strip()
        try:
            parsed = datetime.fromisoformat(text.replace("Z", "+00:00"))
        except ValueError:
            try:
                parsed = parsedate_to_datetime(text)
            except (TypeError, ValueError, OverflowError):
                return text[:100]
    if parsed.tzinfo is None:
        parsed = parsed.replace(tzinfo=UTC)
    return parsed.astimezone(UTC).replace(microsecond=0).isoformat().replace("+00:00", "Z")


def _path_value(value: Any, path: str | None) -> Any:
    if not path:
        return value
    current = value
    for segment in path.split("."):
        if segment == "":
            continue
        if isinstance(current, list):
            try:
                current = current[int(segment)]
            except (ValueError, IndexError):
                return None
        elif isinstance(current, dict):
            current = current.get(segment)
        else:
            return None
        if current is None:
            return None
    return current


class CollectionEngine:
    def __init__(self, database: Database, settings: Settings):
        self.database = database
        self.settings = settings
        self.processor = DocumentProcessor(database, settings)
        self._global_slots = threading.BoundedSemaphore(
            settings.scheduler_max_concurrency
        )
        self._source_slots_guard = threading.Lock()
        self._source_slots: dict[str, tuple[int, threading.BoundedSemaphore]] = {}

    def execute_run(self, run_id: str) -> None:
        run = self.database.get_source_run(run_id)
        if run is None or run["status"] != "queued":
            return
        source = self.database.get_source(run["source_id"])
        if source is None:
            return
        source_slot = self._source_slot(source["id"], source["concurrency_limit"])
        with self._global_slots, source_slot:
            self._execute_with_slot(run_id, source)

    def _execute_with_slot(self, run_id: str, source: Any) -> None:
        request_summary = self._request_summary(source)
        started = self.database.mark_source_run_started(run_id, request_summary)
        if started is None:
            return
        parser_steps = json.loads(started["parser_steps_json"] or "[]")
        parser_steps.append(
            f"任务已领取：优先级 {started['priority']}，第 {started['attempt']}/{started['max_attempts']} 次尝试"
        )
        started_at = time.perf_counter()
        deadline = started_at + max(5, int(started["timeout_seconds"]))
        attempt = int(started["attempt"])
        max_attempts = max(1, int(started["max_attempts"]))
        current_step = "collect"
        try:
            self.database.start_workflow_step(run_id, current_step)
            step_started = time.perf_counter()
            artifacts = self._collect(
                source,
                parser_steps,
                timeout_cap=self._remaining_seconds(deadline),
            )
            self._ensure_within_deadline(deadline)
            self.database.complete_workflow_step(
                run_id,
                current_step,
                f"获取 {len(artifacts)} 条候选内容",
                round((time.perf_counter() - step_started) * 1000),
            )

            current_step = "normalize"
            self.database.start_workflow_step(run_id, current_step)
            step_started = time.perf_counter()
            artifacts = self._normalize_artifacts(artifacts)
            self._ensure_within_deadline(deadline)
            self.database.complete_workflow_step(
                run_id,
                current_step,
                f"标准化 {len(artifacts)} 条内容与元数据",
                round((time.perf_counter() - step_started) * 1000),
            )

            current_step = "deduplicate"
            self.database.start_workflow_step(run_id, current_step)
            step_started = time.perf_counter()
            created = updated = duplicates = 0
            process_document_ids: list[str] = []
            for artifact in artifacts:
                self._ensure_within_deadline(deadline)
                fingerprint_input = (
                    _clean_text(artifact.readable_text).encode("utf-8")
                    or artifact.raw_content
                )
                content_hash = hashlib.sha256(fingerprint_input).hexdigest()
                state, saved_document = self.database.save_collection_document(
                    run_id=run_id,
                    project_id=source["project_id"],
                    source_id=source["id"],
                    canonical_url=artifact.canonical_url,
                    title=artifact.title,
                    published_at=artifact.published_at,
                    language=artifact.language,
                    content_type=artifact.content_type,
                    content_hash=content_hash,
                    raw_content=artifact.raw_content,
                    readable_text=artifact.readable_text,
                    metadata=artifact.metadata,
                    structured_fields=artifact.structured_fields,
                    parser_version=PARSER_VERSION,
                )
                if state == "created":
                    created += 1
                    process_document_ids.append(saved_document["id"])
                elif state == "updated":
                    updated += 1
                    process_document_ids.append(saved_document["id"])
                else:
                    duplicates += 1
            dedupe_summary = f"新增 {created}，更新 {updated}，重复 {duplicates}"
            parser_steps.append(f"指纹去重完成：{dedupe_summary}")
            self.database.complete_workflow_step(
                run_id,
                current_step,
                dedupe_summary,
                round((time.perf_counter() - step_started) * 1000),
            )
            processed = 0
            review_required = 0
            for document_id in process_document_ids:
                self._ensure_within_deadline(deadline)
                processing_row = self.processor.process_document(document_id)
                if processing_row is not None:
                    processed += 1
                    if processing_row["processing_status"] in {"review_required", "failed"}:
                        review_required += 1
            if process_document_ids:
                parser_steps.append(
                    f"清洗与抽取完成：处理 {processed} 条，待复核 {review_required} 条"
                )

            current_step = "quality_gate"
            self.database.start_workflow_step(run_id, current_step)
            step_started = time.perf_counter()
            completeness = self._parser_completeness(artifacts) if artifacts else 100.0
            status_value = "partial" if artifacts and completeness < 70 else "succeeded"
            self._ensure_within_deadline(deadline)
            self.database.complete_workflow_step(
                run_id,
                current_step,
                f"解析完整度 {completeness:.1f}%，状态 {status_value}",
                round((time.perf_counter() - step_started) * 1000),
            )
            elapsed_ms = round((time.perf_counter() - started_at) * 1000)
            self.database.complete_source_run(
                run_id,
                status_value=status_value,
                items_discovered=len(artifacts),
                documents_created=created,
                documents_updated=updated,
                duplicates_skipped=duplicates,
                parser_steps=parser_steps,
                latency_ms=elapsed_ms,
                parser_completeness=completeness,
            )
        except CollectionError as exc:
            self._handle_run_error(
                run_id,
                current_step,
                exc.code,
                str(exc),
                exc.retryable,
                attempt,
                max_attempts,
                int(started["backoff_seconds"]),
                parser_steps,
            )
        except Exception as exc:  # Defensive boundary for connector failures.
            self._handle_run_error(
                run_id,
                current_step,
                "unexpected_error",
                f"{type(exc).__name__}: {exc}",
                True,
                attempt,
                max_attempts,
                int(started["backoff_seconds"]),
                parser_steps,
            )

    def _handle_run_error(
        self,
        run_id: str,
        step_key: str,
        error_type: str,
        error_message: str,
        retryable: bool,
        attempt: int,
        max_attempts: int,
        backoff_seconds: int,
        parser_steps: list[str],
    ) -> None:
        parser_steps.append(f"{error_type}: {error_message}")
        if retryable and attempt < max_attempts:
            delay = min(300, max(1, backoff_seconds) * (2 ** (attempt - 1)))
            parser_steps.append(f"将在 {delay} 秒后执行指数退避重试")
            self.database.schedule_source_run_retry(
                run_id,
                step_key,
                error_type=error_type,
                error_message=error_message,
                delay_seconds=delay,
                parser_steps=parser_steps,
            )
            return
        self.database.fail_workflow_step(run_id, step_key, error_type, error_message)
        self.database.fail_source_run(
            run_id,
            error_type=error_type,
            error_message=error_message,
            parser_steps=parser_steps,
        )

    def _source_slot(
        self, source_id: str, limit: int
    ) -> threading.BoundedSemaphore:
        normalized_limit = max(1, int(limit))
        with self._source_slots_guard:
            current = self._source_slots.get(source_id)
            if current is None or current[0] != normalized_limit:
                current = (
                    normalized_limit,
                    threading.BoundedSemaphore(normalized_limit),
                )
                self._source_slots[source_id] = current
            return current[1]

    @staticmethod
    def _remaining_seconds(deadline: float) -> float:
        return max(1.0, deadline - time.perf_counter())

    @staticmethod
    def _ensure_within_deadline(deadline: float) -> None:
        if time.perf_counter() > deadline:
            raise CollectionError(
                "task_timeout", "工作流超过任务级超时限制", retryable=True
            )

    @staticmethod
    def _normalize_artifacts(
        artifacts: list[CollectedArtifact],
    ) -> list[CollectedArtifact]:
        for artifact in artifacts:
            artifact.title = _clean_text(artifact.title)[:500] or "未命名材料"
            artifact.readable_text = _clean_text(artifact.readable_text)
            artifact.canonical_url = artifact.canonical_url.strip()
        return artifacts

    def _collect(
        self,
        source: Any,
        steps: list[str],
        *,
        timeout_cap: float | None = None,
    ) -> list[CollectedArtifact]:
        source_type = source["source_type"]
        config = json.loads(source["collection_config_json"] or "{}")
        if timeout_cap is not None:
            config["timeout_seconds"] = min(
                float(config.get("timeout_seconds", 20)), timeout_cap
            )
        if source_type in {"webpage", "sitemap"}:
            return self._collect_static_webpage(source, config, steps)
        if source_type == "dynamic_webpage":
            return self._collect_dynamic_webpage(source, config, steps)
        if source_type == "rss":
            return self._collect_rss(source, config, steps)
        if source_type in API_SOURCE_TYPES:
            return self._collect_api(source, config, steps)
        if source_type == "file_upload":
            return self._collect_file(source, config, steps)
        raise CollectionError("unsupported_source_type", f"不支持的来源类型：{source_type}")

    @staticmethod
    def _request_summary(source: Any) -> str:
        method = "FILE" if source["source_type"] == "file_upload" else "GET"
        return f"{method} {source['endpoint']} · {source['source_type']}"

    def _validate_remote_url(self, url: str) -> None:
        parsed = urlparse(url)
        if parsed.scheme not in {"http", "https"} or not parsed.hostname:
            raise CollectionError("invalid_url", "采集地址必须是有效的 HTTP 或 HTTPS URL")
        if parsed.username or parsed.password:
            raise CollectionError("invalid_url", "采集地址不得内嵌用户名或密码")
        if self.settings.collector_allow_private_networks:
            return
        try:
            addresses = {
                item[4][0]
                for item in socket.getaddrinfo(
                    parsed.hostname,
                    parsed.port or (443 if parsed.scheme == "https" else 80),
                )
            }
        except socket.gaierror as exc:
            raise CollectionError(
                "dns_error", f"无法解析来源域名：{parsed.hostname}", retryable=True
            ) from exc
        for address in addresses:
            ip = ipaddress.ip_address(address)
            if (
                ip.is_private
                or ip.is_loopback
                or ip.is_link_local
                or ip.is_multicast
                or ip.is_reserved
                or ip.is_unspecified
            ):
                raise CollectionError("private_network_blocked", "安全策略禁止访问内网或本机地址")

    def _credential_headers(self, source: Any, config: dict[str, Any]) -> dict[str, str]:
        headers = {"User-Agent": self.settings.collector_user_agent}
        accept = config.get("accept")
        if isinstance(accept, str) and len(accept) <= 200:
            headers["Accept"] = accept
        credential_ref = source["credential_ref"]
        if not credential_ref:
            return headers
        try:
            secret_map = json.loads(os.getenv("SOURCE_CREDENTIALS_JSON", "{}"))
        except json.JSONDecodeError:
            secret_map = {}
        secret = secret_map.get(credential_ref)
        if not secret:
            raise CollectionError(
                "credential_unavailable",
                "密钥引用未能从运行环境解析；请在 SOURCE_CREDENTIALS_JSON 中配置该引用",
            )
        header_name = str(config.get("credential_header", "Authorization"))
        if not re.fullmatch(r"[A-Za-z0-9-]{1,64}", header_name):
            raise CollectionError("invalid_credential_header", "凭据请求头名称无效")
        prefix = str(config.get("credential_prefix", "Bearer "))
        headers[header_name] = f"{prefix}{secret}"
        return headers

    def _apply_conditional_headers(
        self, source_id: str, headers: dict[str, str]
    ) -> dict[str, str]:
        latest = self.database.get_latest_source_document(source_id)
        if latest is None:
            return headers
        metadata = json.loads(latest["metadata_json"] or "{}")
        result = dict(headers)
        if metadata.get("etag"):
            result["If-None-Match"] = str(metadata["etag"])
        if metadata.get("last_modified"):
            result["If-Modified-Since"] = str(metadata["last_modified"])
        return result

    def _request_bytes(
        self, url: str, headers: dict[str, str], timeout_seconds: float
    ) -> FetchedResponse:
        current_url = url
        with httpx.Client(timeout=timeout_seconds, follow_redirects=False) as client:
            for _ in range(6):
                self._validate_remote_url(current_url)
                try:
                    with client.stream("GET", current_url, headers=headers) as response:
                        if response.status_code in {301, 302, 303, 307, 308}:
                            location = response.headers.get("location")
                            if not location:
                                raise CollectionError("redirect_error", "重定向响应缺少 Location")
                            current_url = urljoin(current_url, location)
                            continue
                        if response.status_code == 304:
                            return FetchedResponse(
                                current_url,
                                304,
                                b"",
                                response.headers.get("content-type", "application/octet-stream"),
                                dict(response.headers),
                            )
                        response.raise_for_status()
                        declared = int(response.headers.get("content-length", "0") or 0)
                        if declared > self.settings.collector_max_response_bytes:
                            raise CollectionError("response_too_large", "响应超过允许的最大体积")
                        chunks: list[bytes] = []
                        size = 0
                        for chunk in response.iter_bytes():
                            size += len(chunk)
                            if size > self.settings.collector_max_response_bytes:
                                raise CollectionError("response_too_large", "响应超过允许的最大体积")
                            chunks.append(chunk)
                        return FetchedResponse(
                            str(response.url),
                            response.status_code,
                            b"".join(chunks),
                            response.headers.get("content-type", "application/octet-stream"),
                            dict(response.headers),
                        )
                except httpx.HTTPStatusError as exc:
                    retryable = exc.response.status_code >= 500 or exc.response.status_code == 429
                    raise CollectionError(
                        "http_error",
                        f"来源返回 HTTP {exc.response.status_code}",
                        retryable=retryable,
                    ) from exc
                except httpx.RequestError as exc:
                    raise CollectionError(
                        "network_error", f"请求来源失败：{exc}", retryable=True
                    ) from exc
        raise CollectionError("redirect_error", "来源重定向次数超过限制")

    def _check_robots(
        self, url: str, headers: dict[str, str], timeout_seconds: float, steps: list[str]
    ) -> None:
        parsed = urlparse(url)
        robots_url = f"{parsed.scheme}://{parsed.netloc}/robots.txt"
        try:
            response = self._request_bytes(robots_url, headers, timeout_seconds)
        except CollectionError as exc:
            if exc.code == "http_error" and "404" in str(exc):
                steps.append("robots.txt 不存在，按来源登记策略继续")
                return
            if exc.code == "http_error" and any(code in str(exc) for code in ("401", "403")):
                raise CollectionError(
                    "robots_unavailable", "robots.txt 拒绝访问，采集已停止"
                ) from exc
            steps.append("robots.txt 暂时不可达，已依据人工确认记录继续")
            return
        parser = RobotFileParser()
        parser.set_url(robots_url)
        parser.parse(_decode_bytes(response.content, response.content_type).splitlines())
        if not parser.can_fetch(headers["User-Agent"], url):
            raise CollectionError("robots_disallowed", "robots.txt 不允许当前采集器访问该地址")
        steps.append("robots.txt 检查通过")

    def _parse_html_artifact(
        self,
        response: FetchedResponse,
        config: dict[str, Any],
        source_name: str,
    ) -> CollectedArtifact:
        html = _decode_bytes(response.content, response.content_type)
        soup = BeautifulSoup(html, "html.parser")
        for element in soup.select("script, style, noscript, nav, header, footer, aside, form"):
            element.decompose()
        selector = config.get("content_selector")
        content_node = soup.select_one(selector) if isinstance(selector, str) and selector else None
        content_node = content_node or soup.find("main") or soup.find("article") or soup.body or soup
        readable_text = _clean_text(content_node.get_text("\n", strip=True))
        if not readable_text:
            raise CollectionError("empty_content", "网页未提取到可读正文")
        title_node = soup.find("title")
        heading = soup.find("h1")
        title = (
            (heading.get_text(" ", strip=True) if heading else "")
            or (title_node.get_text(" ", strip=True) if title_node else "")
            or source_name
        )
        canonical_node = soup.find("link", rel=lambda value: value and "canonical" in value)
        canonical_url = response.url
        if canonical_node and canonical_node.get("href"):
            canonical_url = urljoin(response.url, str(canonical_node["href"]))
        published = None
        for attrs in (
            {"property": "article:published_time"},
            {"name": "date"},
            {"name": "publishdate"},
        ):
            node = soup.find("meta", attrs=attrs)
            if node and node.get("content"):
                published = _normalise_datetime(node["content"])
                break
        language = soup.html.get("lang") if soup.html else None
        metadata = {
            "http_status": response.status_code,
            "final_url": response.url,
            "etag": response.headers.get("etag"),
            "last_modified": response.headers.get("last-modified"),
            "content_selector": selector,
        }
        return CollectedArtifact(
            canonical_url=canonical_url,
            title=title,
            raw_content=response.content,
            readable_text=readable_text,
            content_type=response.content_type,
            published_at=published,
            language=str(language)[:20] if language else None,
            metadata=metadata,
            structured_fields={"title": title, "published_at": published},
        )

    def _collect_static_webpage(
        self, source: Any, config: dict[str, Any], steps: list[str]
    ) -> list[CollectedArtifact]:
        timeout = float(config.get("timeout_seconds", 20))
        base_headers = self._credential_headers(source, config)
        headers = self._apply_conditional_headers(
            source["id"], base_headers
        )
        self._check_robots(source["endpoint"], base_headers, timeout, steps)
        response = self._request_bytes(source["endpoint"], headers, timeout)
        steps.append(f"静态页面请求完成：HTTP {response.status_code}")
        if response.status_code == 304:
            steps.append("条件请求命中 304，来源内容未变化")
            return []
        if source["source_type"] == "sitemap":
            return self._parse_sitemap(response, source, config, steps)
        artifact = self._parse_html_artifact(response, config, source["name"])
        steps.append("正文、标题、时间和规范 URL 抽取完成")
        return [artifact]

    def _parse_sitemap(
        self,
        response: FetchedResponse,
        source: Any,
        config: dict[str, Any],
        steps: list[str],
    ) -> list[CollectedArtifact]:
        soup = BeautifulSoup(response.content, "xml")
        max_items = int(config.get("max_items", 50))
        urls = [node.get_text(strip=True) for node in soup.find_all("loc")][:max_items]
        if not urls:
            raise CollectionError("invalid_sitemap", "站点地图中没有可用 URL")
        text = "\n".join(urls)
        steps.append(f"站点地图解析完成：发现 {len(urls)} 个 URL")
        return [
            CollectedArtifact(
                canonical_url=response.url,
                title=f"{source['name']} · 站点地图",
                raw_content=response.content,
                readable_text=text,
                content_type=response.content_type,
                metadata={"url_count": len(urls)},
                structured_fields={"urls": urls},
            )
        ]

    def _collect_dynamic_webpage(
        self, source: Any, config: dict[str, Any], steps: list[str]
    ) -> list[CollectedArtifact]:
        timeout = float(config.get("timeout_seconds", 30))
        headers = self._credential_headers(source, config)
        self._validate_remote_url(source["endpoint"])
        self._check_robots(source["endpoint"], headers, timeout, steps)
        try:
            from playwright.sync_api import Error as PlaywrightError
            from playwright.sync_api import sync_playwright
        except ImportError as exc:
            raise CollectionError(
                "dynamic_browser_unavailable",
                "动态网页连接器未安装 Playwright",
            ) from exc
        try:
            with sync_playwright() as playwright:
                channel = str(config.get("browser_channel", "msedge"))
                try:
                    browser = playwright.chromium.launch(headless=True, channel=channel)
                except PlaywrightError:
                    browser = playwright.chromium.launch(headless=True)
                context = browser.new_context(
                    user_agent=self.settings.collector_user_agent,
                    extra_http_headers={key: value for key, value in headers.items() if key != "User-Agent"},
                )
                page = context.new_page()

                def route_guard(route: Any) -> None:
                    request_url = route.request.url
                    if urlparse(request_url).scheme in {"data", "blob"}:
                        route.continue_()
                        return
                    try:
                        self._validate_remote_url(request_url)
                    except CollectionError:
                        route.abort()
                        return
                    route.continue_()

                page.route("**/*", route_guard)
                page.goto(
                    source["endpoint"],
                    wait_until="domcontentloaded",
                    timeout=round(timeout * 1000),
                )
                selector = config.get("wait_selector")
                if isinstance(selector, str) and selector:
                    page.wait_for_selector(selector, timeout=round(timeout * 1000))
                wait_ms = min(15000, max(0, int(config.get("wait_ms", 800))))
                if wait_ms:
                    page.wait_for_timeout(wait_ms)
                html = page.content().encode("utf-8")
                final_url = page.url
                browser.close()
        except PlaywrightError as exc:
            raise CollectionError(
                "dynamic_render_error", f"动态页面渲染失败：{exc}", retryable=True
            ) from exc
        steps.append("浏览器渲染与脚本执行完成")
        response = FetchedResponse(
            final_url,
            200,
            html,
            "text/html; charset=utf-8",
            {},
        )
        artifact = self._parse_html_artifact(response, config, source["name"])
        artifact.metadata["rendered"] = True
        artifact.metadata["wait_selector"] = config.get("wait_selector")
        steps.append("动态 DOM 正文抽取完成")
        return [artifact]

    def _collect_rss(
        self, source: Any, config: dict[str, Any], steps: list[str]
    ) -> list[CollectedArtifact]:
        timeout = float(config.get("timeout_seconds", 20))
        headers = self._apply_conditional_headers(
            source["id"], self._credential_headers(source, config)
        )
        response = self._request_bytes(source["endpoint"], headers, timeout)
        if response.status_code == 304:
            steps.append("条件请求命中 304，RSS/Atom 未更新")
            return []
        parsed = feedparser.parse(response.content)
        if parsed.bozo and not parsed.entries:
            raise CollectionError("invalid_feed", f"RSS/Atom 解析失败：{parsed.bozo_exception}")
        max_items = int(config.get("max_items", 50))
        artifacts: list[CollectedArtifact] = []
        for index, entry in enumerate(parsed.entries[:max_items]):
            title = str(entry.get("title") or f"{source['name']} 第 {index + 1} 条")
            link = str(entry.get("link") or f"{response.url}#entry-{index + 1}")
            html_content = ""
            if entry.get("content"):
                html_content = str(entry.content[0].get("value", ""))
            html_content = html_content or str(entry.get("summary") or entry.get("description") or title)
            readable = _clean_text(BeautifulSoup(html_content, "html.parser").get_text("\n", strip=True))
            published = _normalise_datetime(entry.get("published") or entry.get("updated"))
            raw_entry = json.dumps(dict(entry), ensure_ascii=False, default=str).encode("utf-8")
            artifacts.append(
                CollectedArtifact(
                    canonical_url=link,
                    title=title,
                    raw_content=raw_entry,
                    readable_text=readable,
                    content_type="application/feed+json",
                    published_at=published,
                    language=parsed.feed.get("language"),
                    metadata={
                        "feed_url": response.url,
                        "feed_title": parsed.feed.get("title"),
                        "entry_id": entry.get("id"),
                        "etag": response.headers.get("etag"),
                        "last_modified": response.headers.get("last-modified"),
                    },
                    structured_fields={
                        "title": title,
                        "url": link,
                        "published_at": published,
                    },
                )
            )
        if not artifacts:
            raise CollectionError("empty_feed", "RSS/Atom 中没有可采集条目")
        steps.append(f"RSS/Atom 解析完成：发现 {len(artifacts)} 条")
        return artifacts

    def _collect_api(
        self, source: Any, config: dict[str, Any], steps: list[str]
    ) -> list[CollectedArtifact]:
        timeout = float(config.get("timeout_seconds", 20))
        headers = self._apply_conditional_headers(
            source["id"],
            self._credential_headers(source, {**config, "accept": "application/json"}),
        )
        response = self._request_bytes(source["endpoint"], headers, timeout)
        if response.status_code == 304:
            steps.append("条件请求命中 304，API 数据未更新")
            return []
        try:
            payload = json.loads(_decode_bytes(response.content, response.content_type))
        except json.JSONDecodeError as exc:
            raise CollectionError("invalid_json", "API 响应不是有效 JSON") from exc
        items_path = config.get("items_path")
        items = _path_value(payload, str(items_path)) if items_path else None
        if items is None:
            if isinstance(payload, list):
                items = payload
            elif isinstance(payload, dict):
                items = next(
                    (payload[key] for key in ("items", "data", "results") if isinstance(payload.get(key), list)),
                    [payload],
                )
        if not isinstance(items, list):
            items = [items]
        max_items = int(config.get("max_items", 50))
        artifacts: list[CollectedArtifact] = []
        field_mapping = config.get("field_mapping") if isinstance(config.get("field_mapping"), dict) else {}
        for index, item in enumerate(items[:max_items]):
            if not isinstance(item, dict):
                item = {"value": item}
            title = self._first_value(
                item,
                [field_mapping.get("title"), config.get("title_path"), "title", "name", "label", "id"],
            ) or f"{source['name']} 第 {index + 1} 条"
            body_value = self._first_value(
                item,
                [field_mapping.get("body"), config.get("content_path"), "body", "content", "description", "summary"],
            )
            readable = _clean_text(
                str(body_value) if body_value is not None else json.dumps(item, ensure_ascii=False, default=str)
            )
            url_value = self._first_value(
                item,
                [field_mapping.get("url"), config.get("url_path"), "url", "html_url", "link", "id"],
            )
            canonical_url = str(url_value) if url_value else f"{response.url}#item-{index + 1}"
            if not urlparse(canonical_url).scheme:
                canonical_url = urljoin(response.url, canonical_url)
            published_value = self._first_value(
                item,
                [field_mapping.get("published_at"), config.get("published_path"), "published_at", "published", "created_at", "updated_at"],
            )
            structured = {}
            for field_name in json.loads(source["fields_json"] or "[]"):
                mapped_path = field_mapping.get(field_name, field_name)
                value = _path_value(item, str(mapped_path))
                if value is not None:
                    structured[field_name] = value
            artifacts.append(
                CollectedArtifact(
                    canonical_url=canonical_url,
                    title=str(title),
                    raw_content=json.dumps(item, ensure_ascii=False, default=str).encode("utf-8"),
                    readable_text=readable,
                    content_type="application/json",
                    published_at=_normalise_datetime(published_value),
                    language=str(config.get("language")) if config.get("language") else None,
                    metadata={
                        "api_url": response.url,
                        "item_index": index,
                        "etag": response.headers.get("etag"),
                        "last_modified": response.headers.get("last-modified"),
                    },
                    structured_fields=structured or item,
                )
            )
        if not artifacts:
            raise CollectionError("empty_api_response", "API 响应中没有可采集记录")
        steps.append(f"JSON 字段映射完成：发现 {len(artifacts)} 条")
        return artifacts

    @staticmethod
    def _first_value(item: dict[str, Any], paths: list[Any]) -> Any:
        for path in paths:
            if not path:
                continue
            value = _path_value(item, str(path))
            if value is not None and value != "":
                return value
        return None

    def _collect_file(
        self, source: Any, config: dict[str, Any], steps: list[str]
    ) -> list[CollectedArtifact]:
        storage_value = config.get("_storage_path")
        if not storage_value:
            raise CollectionError("file_missing", "文件来源尚未上传原始文件")
        path = Path(str(storage_value)).resolve()
        uploads_root = (self.database.path.parent / "uploads").resolve()
        if uploads_root not in path.parents:
            raise CollectionError("file_path_blocked", "文件路径不在受管上传目录中")
        if not path.is_file():
            raise CollectionError("file_missing", "上传文件不存在或已被移除")
        raw = path.read_bytes()
        if len(raw) > self.settings.collector_max_response_bytes:
            raise CollectionError("file_too_large", "上传文件超过允许的最大体积")
        suffix = path.suffix.lower()
        metadata: dict[str, Any] = {"filename": path.name, "size_bytes": len(raw)}
        structured: dict[str, Any] = {}
        language = None
        if suffix in {".txt", ".md", ".log"}:
            text = _decode_bytes(raw)
            content_type = "text/plain"
        elif suffix in {".html", ".htm"}:
            response = FetchedResponse(source["endpoint"], 200, raw, "text/html", {})
            artifact = self._parse_html_artifact(response, config, source["name"])
            artifact.metadata.update(metadata)
            steps.append("HTML 文件正文抽取完成")
            return [artifact]
        elif suffix == ".json":
            try:
                payload = json.loads(_decode_bytes(raw))
            except json.JSONDecodeError as exc:
                raise CollectionError("invalid_json_file", "JSON 文件格式无效") from exc
            text = json.dumps(payload, ensure_ascii=False, indent=2)
            structured = payload if isinstance(payload, dict) else {"items": payload}
            content_type = "application/json"
        elif suffix == ".csv":
            decoded = _decode_bytes(raw)
            rows = list(csv.DictReader(io.StringIO(decoded)))
            text = "\n".join(" | ".join(f"{k}: {v}" for k, v in row.items()) for row in rows)
            structured = {"columns": list(rows[0]) if rows else [], "row_count": len(rows)}
            metadata["row_count"] = len(rows)
            content_type = "text/csv"
        elif suffix in {".xml", ".rss", ".atom"}:
            soup = BeautifulSoup(raw, "xml")
            text = _clean_text(soup.get_text("\n", strip=True))
            content_type = "application/xml"
        elif suffix == ".pdf":
            try:
                from pypdf import PdfReader
            except ImportError as exc:
                raise CollectionError("file_parser_unavailable", "PDF 解析器未安装") from exc
            reader = PdfReader(io.BytesIO(raw))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
            metadata["page_count"] = len(reader.pages)
            content_type = "application/pdf"
        elif suffix == ".docx":
            try:
                from docx import Document
            except ImportError as exc:
                raise CollectionError("file_parser_unavailable", "DOCX 解析器未安装") from exc
            document = Document(io.BytesIO(raw))
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            metadata["paragraph_count"] = len(document.paragraphs)
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        elif suffix == ".xlsx":
            try:
                from openpyxl import load_workbook
            except ImportError as exc:
                raise CollectionError("file_parser_unavailable", "XLSX 解析器未安装") from exc
            workbook = load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
            lines: list[str] = []
            for sheet in workbook.worksheets:
                lines.append(f"[{sheet.title}]")
                for row in sheet.iter_rows(values_only=True):
                    lines.append(" | ".join("" if value is None else str(value) for value in row))
            text = "\n".join(lines)
            metadata["sheet_count"] = len(workbook.worksheets)
            content_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        elif suffix in {".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff", ".bmp"}:
            content_type = {
                ".png": "image/png",
                ".jpg": "image/jpeg",
                ".jpeg": "image/jpeg",
                ".webp": "image/webp",
                ".tif": "image/tiff",
                ".tiff": "image/tiff",
                ".bmp": "image/bmp",
            }[suffix]
            metadata["ocr_required"] = True
            steps.append(f"{suffix} 图像已进入 OCR 处理队列")
            return [
                CollectedArtifact(
                    canonical_url=source["endpoint"],
                    title=path.name,
                    raw_content=raw,
                    readable_text="",
                    content_type=content_type,
                    language=None,
                    metadata=metadata,
                    structured_fields={"ocr_required": True},
                )
            ]
        else:
            raise CollectionError(
                "unsupported_file_type",
                "仅支持 TXT、Markdown、HTML、JSON、CSV、XML、PDF、DOCX、XLSX 和常见图片文件",
            )
        text = _clean_text(text)
        if not text:
            raise CollectionError("empty_file", "文件未提取到可读内容")
        steps.append(f"{suffix or '未知类型'} 文件解析完成")
        return [
            CollectedArtifact(
                canonical_url=source["endpoint"],
                title=path.name,
                raw_content=raw,
                readable_text=text,
                content_type=content_type,
                language=language,
                metadata=metadata,
                structured_fields=structured,
            )
        ]

    @staticmethod
    def _parser_completeness(artifacts: list[CollectedArtifact]) -> float:
        if not artifacts:
            return 0.0
        scores = []
        for item in artifacts:
            present = sum(
                bool(value)
                for value in (item.title, item.canonical_url, item.readable_text, item.content_type)
            )
            scores.append(present / 4 * 100)
        return round(sum(scores) / len(scores), 1)


class CollectionScheduler:
    def __init__(
        self,
        database: Database,
        engine: CollectionEngine,
        poll_seconds: int,
        dispatch_batch: int = 20,
    ):
        self.database = database
        self.engine = engine
        self.poll_seconds = poll_seconds
        self.dispatch_batch = max(1, dispatch_batch)
        self._task: asyncio.Task[None] | None = None
        self._run_tasks: set[asyncio.Task[None]] = set()
        self._inflight_ids: set[str] = set()

    def start(self) -> None:
        if self._task is None:
            self._task = asyncio.create_task(self._loop(), name="collection-scheduler")

    async def stop(self) -> None:
        if self._task is not None:
            self._task.cancel()
            try:
                await self._task
            except asyncio.CancelledError:
                pass
            self._task = None
        if self._run_tasks:
            await asyncio.gather(*self._run_tasks, return_exceptions=True)

    async def _loop(self) -> None:
        while True:
            self.database.queue_due_source_runs(self.dispatch_batch)
            capacity = max(0, self.dispatch_batch - len(self._run_tasks))
            for run_id in self.database.list_dispatchable_run_ids(capacity):
                if run_id in self._inflight_ids:
                    continue
                self._inflight_ids.add(run_id)
                task = asyncio.create_task(
                    asyncio.to_thread(self.engine.execute_run, run_id)
                )
                self._run_tasks.add(task)
                task.add_done_callback(
                    lambda finished, item_id=run_id: self._run_finished(
                        finished, item_id
                    )
                )
            await asyncio.sleep(self.poll_seconds)

    def _run_finished(self, task: asyncio.Task[None], run_id: str) -> None:
        self._run_tasks.discard(task)
        self._inflight_ids.discard(run_id)
